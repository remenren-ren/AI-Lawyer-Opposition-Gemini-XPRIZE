import html
import io
from datetime import datetime, timezone
from typing import Any, Dict, Iterable, List, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


INK = colors.HexColor("#172033")
MUTED = colors.HexColor("#687386")
BRAND = colors.HexColor("#3157D5")
LINE = colors.HexColor("#DCE2EA")
SOFT = colors.HexColor("#F4F6FA")
WARNING = colors.HexColor("#FFF8E7")

REPORT_TYPES = {
    "lawyer_review_pack": "Lawyer Review Pack",
    "client_intake_handoff": "Client Intake Handoff Report",
    "chronology_evidence_index": "Chronology and Evidence Index",
    "professional_formatted_report": "Professional Formatted Report",
    "compliance_review_matrix": "Compliance Review Matrix",
}

SECTION_ORDER: List[Tuple[str, str]] = [
    ("executive_summary", "Executive Summary"),
    ("background", "Background and Request"),
    ("involved_people", "People and Organisations"),
    ("key_dates", "Key Dates"),
    ("chronology", "Chronology"),
    ("client_position", "Client Position or Requested Deliverables"),
    ("other_position", "Other-Side Position or Constraints"),
    ("issues_for_review", "Issues for Lawyer Review"),
    ("evidence_index", "Evidence Index"),
    ("missing_information", "Missing Information and Material"),
    ("requested_deliverables", "Requested Deliverables"),
    ("lawyer_review_tasks", "Lawyer Review Tasks"),
    ("scope_and_limits", "Scope and Limits"),
]


def _text(value: Any) -> str:
    return str(value or "").strip()


def _list(value: Any) -> List[Any]:
    if value is None or value == "":
        return []
    if isinstance(value, list):
        return list(value)
    if isinstance(value, tuple):
        return list(value)
    return [value]


def normalise_output_requirements(value: Any) -> Dict[str, Any]:
    raw = value if isinstance(value, dict) else {}
    report_type = _text(raw.get("report_type")).lower() or "lawyer_review_pack"
    if report_type not in REPORT_TYPES:
        report_type = "lawyer_review_pack"
    desired_format = _text(raw.get("desired_format")).lower() or "docx_pdf"
    if desired_format not in {"docx_pdf", "docx", "pdf", "screen"}:
        desired_format = "docx_pdf"
    return {
        "report_type": report_type,
        "report_type_label": REPORT_TYPES[report_type],
        "jurisdiction_or_body": _text(raw.get("jurisdiction_or_body"))[:240],
        "desired_format": desired_format,
        "formatting_rules": _text(raw.get("formatting_rules"))[:2000],
        "include_source_references": bool(raw.get("include_source_references", True)),
        "include_contents": bool(raw.get("include_contents", True)),
        "include_page_numbers": bool(raw.get("include_page_numbers", True)),
        "include_evidence_index": bool(raw.get("include_evidence_index", True)),
    }


def normalise_professional_report(
    report: Any,
    output_requirements: Any = None,
) -> Dict[str, Any]:
    value = dict(report) if isinstance(report, dict) else {}
    requirements = normalise_output_requirements(
        output_requirements or value.get("output_requirements")
    )
    value["output_requirements"] = requirements
    value["request_type"] = _text(value.get("request_type")) or "legal matter or support request"
    value["case_name"] = _text(value.get("case_name")) or "Client matter"
    value["jurisdiction"] = (
        requirements.get("jurisdiction_or_body")
        or _text(value.get("jurisdiction"))
        or "Not specified"
    )
    value["executive_summary"] = (
        _text(value.get("executive_summary"))
        or _text(value.get("background"))
        or "The supplied material has been organised for professional review."
    )
    value["background"] = _text(value.get("background"))
    value["involved_people"] = _list(value.get("involved_people") or value.get("people_and_organisations"))
    value["key_dates"] = _list(value.get("key_dates"))
    value["chronology"] = _list(value.get("chronology"))
    value["client_position"] = _list(value.get("client_position") or value.get("pos_args"))
    value["other_position"] = _list(value.get("other_position") or value.get("neg_args"))
    value["issues_for_review"] = _list(value.get("issues_for_review"))
    value["evidence_index"] = _list(value.get("evidence_index") or value.get("pos_ev"))
    other_evidence = _list(value.get("neg_ev"))
    if other_evidence:
        value["evidence_index"].extend(other_evidence)
    value["missing_information"] = _list(
        value.get("missing_information") or value.get("missing_details")
    )
    value["requested_deliverables"] = _list(value.get("requested_deliverables"))
    value["lawyer_review_tasks"] = _list(value.get("lawyer_review_tasks") or value.get("next_review_tasks"))
    value["scope_and_limits"] = _list(
        value.get("scope_and_limits") or value.get("analysis_limits") or value.get("scope_notice")
    )
    metadata = value.get("analysis_metadata") if isinstance(value.get("analysis_metadata"), dict) else {}
    value["analysis_metadata"] = {
        "provider_display_name": _text(metadata.get("provider_display_name") or value.get("provider")) or "Not recorded",
        "model_name": _text(metadata.get("model_name")) or "Not recorded",
        "engine_source": _text(metadata.get("engine_source") or value.get("engine_source")) or "Not recorded",
        "generated_at_utc": _text(metadata.get("generated_at_utc")) or datetime.now(timezone.utc).isoformat(timespec="seconds"),
    }
    return value


def _safe(value: Any) -> str:
    return html.escape(_text(value)).replace("\n", "<br/>")


def _iter_rows(value: Any) -> Iterable[Tuple[str, str]]:
    if isinstance(value, dict):
        for key, item in value.items():
            yield str(key).replace("_", " ").title(), _text(item)
    elif isinstance(value, list):
        for index, item in enumerate(value, 1):
            if isinstance(item, dict):
                parts = [f"{str(k).replace('_', ' ').title()}: {_text(v)}" for k, v in item.items() if _text(v)]
                yield str(index), " | ".join(parts)
            else:
                yield str(index), _text(item)
    elif _text(value):
        yield "", _text(value)


def _has_content(value: Any) -> bool:
    if isinstance(value, (list, dict)):
        return bool(value)
    return bool(_text(value))


def _presentation_value(value: Any, include_source_references: bool) -> Any:
    if include_source_references:
        return value
    source_keys = {
        "citation", "page", "page_number", "source", "source_file",
        "source_location", "source_reference",
    }
    if isinstance(value, dict):
        return {
            key: _presentation_value(item, False)
            for key, item in value.items()
            if str(key).strip().lower() not in source_keys
        }
    if isinstance(value, list):
        return [_presentation_value(item, False) for item in value]
    return value


def _pdf_page(canvas, document, firm_name: str, include_page_numbers: bool) -> None:
    canvas.saveState()
    width, height = A4
    canvas.setStrokeColor(LINE)
    canvas.line(18 * mm, height - 13 * mm, width - 18 * mm, height - 13 * mm)
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(MUTED)
    canvas.drawString(18 * mm, height - 10 * mm, firm_name[:70])
    if include_page_numbers:
        canvas.drawRightString(width - 18 * mm, 10 * mm, f"Page {document.page}")
    canvas.drawString(18 * mm, 10 * mm, "AI-assisted lawyer reference material - professional review required")
    canvas.restoreState()


def build_professional_pdf(record: Dict[str, Any], firm_name: str) -> bytes:
    report = normalise_professional_report(
        record.get("free_report") or {}, record.get("output_requirements")
    )
    requirements = report["output_requirements"]
    metadata = report["analysis_metadata"]
    output = io.BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=20 * mm,
        bottomMargin=18 * mm,
        title=requirements["report_type_label"],
        author=firm_name,
    )
    base = getSampleStyleSheet()
    title = ParagraphStyle("ProfessionalTitle", parent=base["Title"], fontName="Helvetica-Bold", fontSize=22, leading=27, textColor=INK, alignment=TA_CENTER, spaceAfter=8)
    subtitle = ParagraphStyle("ProfessionalSubtitle", parent=base["Normal"], fontSize=9.5, leading=14, textColor=MUTED, alignment=TA_CENTER, spaceAfter=14)
    heading = ParagraphStyle("ProfessionalHeading", parent=base["Heading2"], fontName="Helvetica-Bold", fontSize=15, leading=19, textColor=BRAND, spaceBefore=10, spaceAfter=7)
    body = ParagraphStyle("ProfessionalBody", parent=base["BodyText"], fontSize=10.2, leading=14.5, textColor=INK, spaceAfter=6)
    notice = ParagraphStyle("ProfessionalNotice", parent=body, fontSize=9, leading=13, textColor=INK, backColor=WARNING, borderColor=colors.HexColor("#E1BF69"), borderWidth=0.5, borderPadding=8, spaceBefore=14)

    story: List[Any] = [
        Paragraph(_safe(firm_name), subtitle),
        Paragraph(_safe(requirements["report_type_label"]), title),
        Paragraph(
            f"<b>Matter:</b> {_safe(report['case_name'])}<br/>"
            f"<b>Jurisdiction/body:</b> {_safe(report['jurisdiction'])}<br/>"
            f"<b>AI provider:</b> {_safe(metadata['provider_display_name'])}<br/>"
            f"<b>Model:</b> {_safe(metadata['model_name'])}<br/>"
            f"<b>Generated:</b> {_safe(metadata['generated_at_utc'])}",
            subtitle,
        ),
    ]
    if requirements.get("include_contents"):
        visible = [label for key, label in SECTION_ORDER if _has_content(report.get(key))]
        story.extend([
            Paragraph("Report contents", heading),
            Paragraph("<br/>".join(f"{index}. {_safe(label)}" for index, label in enumerate(visible, 1)), body),
            Spacer(1, 4 * mm),
        ])
    for key, label in SECTION_ORDER:
        value = _presentation_value(
            report.get(key), requirements.get("include_source_references", True)
        )
        if key == "evidence_index" and not requirements.get("include_evidence_index"):
            continue
        if not _has_content(value):
            continue
        blocks: List[Any] = [Paragraph(_safe(label), heading)]
        rows = list(_iter_rows(value))
        if len(rows) == 1 and not rows[0][0]:
            blocks.append(Paragraph(_safe(rows[0][1]), body))
        else:
            data = [[Paragraph("No.", body), Paragraph("Details", body)]]
            data.extend([[Paragraph(_safe(left), body), Paragraph(_safe(right), body)] for left, right in rows if right])
            table = Table(data, colWidths=[20 * mm, 148 * mm], repeatRows=1, hAlign="LEFT")
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), SOFT),
                ("TEXTCOLOR", (0, 0), (-1, 0), INK),
                ("GRID", (0, 0), (-1, -1), 0.4, LINE),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]))
            blocks.append(table)
        story.append(KeepTogether(blocks) if len(rows) < 6 else blocks[0])
        if len(rows) >= 6:
            story.extend(blocks[1:])
    if requirements.get("formatting_rules"):
        story.extend([
            Paragraph("Requested Formatting Instructions", heading),
            Paragraph(_safe(requirements["formatting_rules"]), body),
        ])
    story.extend([
        PageBreak(),
        Paragraph("IMPORTANT PROFESSIONAL REVIEW NOTICE", heading),
        Paragraph(
            "This document is AI-assisted lawyer reference material generated from information supplied by the user. "
            "It is not legal advice, a final legal opinion, or proof of compliance with any court or regulator's filing rules. "
            "A qualified lawyer must verify the facts, source documents, law, jurisdiction, deadlines, citations, procedural "
            "requirements and final wording. Any selected formatting template controls presentation only unless the relevant "
            "official rule set and version have been independently verified.",
            notice,
        ),
    ])
    page_callback = lambda canvas, doc: _pdf_page(
        canvas, doc, firm_name, requirements.get("include_page_numbers", True)
    )
    document.build(story, onFirstPage=page_callback, onLaterPages=page_callback)
    return output.getvalue()


def _docx_add_value(document: Any, value: Any) -> None:
    rows = list(_iter_rows(value))
    if len(rows) == 1 and not rows[0][0]:
        document.add_paragraph(rows[0][1])
        return
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "No."
    table.rows[0].cells[1].text = "Details"
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right


def _docx_add_page_number(paragraph: Any) -> None:
    paragraph.add_run("  |  Page ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def build_professional_docx(record: Dict[str, Any], firm_name: str) -> bytes:
    report = normalise_professional_report(
        record.get("free_report") or {}, record.get("output_requirements")
    )
    requirements = report["output_requirements"]
    metadata = report["analysis_metadata"]
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    header = section.header.paragraphs[0]
    header.text = firm_name
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = section.footer.paragraphs[0]
    footer.text = "AI-assisted lawyer reference material - professional review required"
    if requirements.get("include_page_numbers"):
        _docx_add_page_number(footer)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(requirements["report_type_label"])
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(23, 32, 51)
    meta = document.add_table(rows=0, cols=2)
    meta.style = "Table Grid"
    for label, value in (
        ("Matter", report["case_name"]),
        ("Jurisdiction/body", report["jurisdiction"]),
        ("AI provider", metadata["provider_display_name"]),
        ("Model", metadata["model_name"]),
        ("Generated", metadata["generated_at_utc"]),
    ):
        cells = meta.add_row().cells
        cells[0].text = label
        cells[1].text = value
    if requirements.get("include_contents"):
        document.add_heading("Report contents", level=1)
        for key, label in SECTION_ORDER:
            if key == "evidence_index" and not requirements.get("include_evidence_index"):
                continue
            if _has_content(report.get(key)):
                document.add_paragraph(label, style="List Number")
    for key, label in SECTION_ORDER:
        value = _presentation_value(
            report.get(key), requirements.get("include_source_references", True)
        )
        if key == "evidence_index" and not requirements.get("include_evidence_index"):
            continue
        if not _has_content(value):
            continue
        document.add_heading(label, level=1)
        _docx_add_value(document, value)
    if requirements.get("formatting_rules"):
        document.add_heading("Requested Formatting Instructions", level=1)
        document.add_paragraph(requirements["formatting_rules"])
    document.add_page_break()
    document.add_heading("Important Professional Review Notice", level=1)
    document.add_paragraph(
        "This document is AI-assisted lawyer reference material generated from information supplied by the user. "
        "It is not legal advice, a final legal opinion, or proof of compliance with any court or regulator's filing rules. "
        "A qualified lawyer must verify the facts, source documents, law, jurisdiction, deadlines, citations, procedural "
        "requirements and final wording. A selected template controls presentation only unless the relevant official rule "
        "set and version have been independently verified."
    )
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()
