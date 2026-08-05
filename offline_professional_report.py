"""Professional lawyer-work-product exports for the offline application.

The module formats supplied application state only. It does not infer facts,
law, dates, authorities, or filing compliance.
"""

import html
import io
import json
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Iterable, List, Tuple


VENDOR_DIR = Path(__file__).resolve().parent / ".vendor"
if VENDOR_DIR.exists() and str(VENDOR_DIR) not in sys.path:
    sys.path.insert(0, str(VENDOR_DIR))


REPORT_TYPES = {
    "lawyer_working_paper": "Lawyer Working Paper",
    "whole_case_analysis": "Whole-Case Analysis Report",
    "chronology_evidence": "Chronology and Evidence Index",
    "weakness_evidence_matrix": "Weakness and Evidence Matrix",
    "client_readable": "Client-Readable Review Report",
}

SECTION_ORDER: List[Tuple[str, str]] = [
    ("executive_summary", "Executive Summary"),
    ("case_profile", "Matter Profile"),
    ("chronology", "Chronology"),
    ("positions", "Positions and Contentions"),
    ("evidence_index", "Evidence Index"),
    ("issues_and_weaknesses", "Issues and Weaknesses for Review"),
    ("dimension_analysis", "Analysis by Selected Dimension"),
    ("missing_information", "Missing Information and Evidence Gaps"),
    ("risk_and_confidence", "Risk and Confidence Notes"),
    ("lawyer_verification_tasks", "Lawyer Verification Tasks"),
    ("scope_and_limits", "Scope and Limits"),
]

REPORT_SECTION_FILTERS = {
    "chronology_evidence": {
        "executive_summary", "case_profile", "chronology", "evidence_index",
        "missing_information", "lawyer_verification_tasks", "scope_and_limits",
    },
    "weakness_evidence_matrix": {
        "executive_summary", "case_profile", "evidence_index", "issues_and_weaknesses",
        "dimension_analysis", "missing_information", "risk_and_confidence",
        "lawyer_verification_tasks", "scope_and_limits",
    },
    "client_readable": {
        "executive_summary", "case_profile", "chronology", "positions", "evidence_index",
        "issues_and_weaknesses", "missing_information", "scope_and_limits",
    },
}


def _text(value: Any) -> str:
    return str(value or "").strip()


def _list(value: Any) -> List[Any]:
    if value is None or value == "":
        return []
    if isinstance(value, list):
        return list(value)
    if isinstance(value, tuple):
        return list(value)
    return [value]


def normalise_professional_record(record: Any) -> Dict[str, Any]:
    value = dict(record) if isinstance(record, dict) else {}
    report_type = _text(value.get("report_type")) or "lawyer_working_paper"
    if report_type not in REPORT_TYPES:
        report_type = "lawyer_working_paper"
    value["report_type"] = report_type
    value["report_title"] = REPORT_TYPES[report_type]
    value["matter_name"] = _text(value.get("matter_name")) or "Current matter"
    value["jurisdiction"] = _text(value.get("jurisdiction")) or "Not specified"
    value["generated_at_utc"] = _text(value.get("generated_at_utc")) or datetime.now(timezone.utc).isoformat(timespec="seconds")
    value["engine_metadata"] = dict(value.get("engine_metadata") or {})
    value["engine_metadata"].setdefault("provider", "Offline local workflow")
    value["engine_metadata"].setdefault("model", "No external model recorded")
    value["engine_metadata"].setdefault("source", "Local application state")
    for key, _ in SECTION_ORDER:
        if key in {"executive_summary"}:
            value[key] = _text(value.get(key))
        elif key in {"case_profile", "positions"}:
            value[key] = dict(value.get(key) or {})
        else:
            value[key] = _list(value.get(key))
    value["include_contents"] = bool(value.get("include_contents", True))
    value["include_page_numbers"] = bool(value.get("include_page_numbers", True))
    value["include_source_references"] = bool(value.get("include_source_references", True))
    value["include_evidence_index"] = bool(value.get("include_evidence_index", True))
    return value


def _visible_sections(record: Dict[str, Any]) -> Iterable[Tuple[str, str, Any]]:
    allowed = REPORT_SECTION_FILTERS.get(record.get("report_type"))
    for key, label in SECTION_ORDER:
        if allowed is not None and key not in allowed:
            continue
        if key == "evidence_index" and not record.get("include_evidence_index"):
            continue
        value = record.get(key)
        if isinstance(value, (list, dict)) and not value:
            continue
        if not isinstance(value, (list, dict)) and not _text(value):
            continue
        yield key, label, value


def _presentation_value(value: Any, include_source_references: bool) -> Any:
    if include_source_references:
        return value
    hidden = {"citation", "page", "page_number", "source", "source_file", "source_location", "source_reference"}
    if isinstance(value, dict):
        return {
            key: _presentation_value(item, False)
            for key, item in value.items()
            if str(key).strip().lower() not in hidden
        }
    if isinstance(value, list):
        return [_presentation_value(item, False) for item in value]
    return value


def _rows(value: Any) -> Iterable[Tuple[str, str]]:
    if isinstance(value, dict):
        for key, item in value.items():
            yield str(key).replace("_", " ").title(), _text(item)
    elif isinstance(value, list):
        for index, item in enumerate(value, 1):
            if isinstance(item, dict):
                details = " | ".join(
                    f"{str(key).replace('_', ' ').title()}: {_text(item_value)}"
                    for key, item_value in item.items() if _text(item_value)
                )
                yield str(index), details
            else:
                yield str(index), _text(item)
    elif _text(value):
        yield "", _text(value)


def build_professional_markdown(record: Dict[str, Any]) -> str:
    value = normalise_professional_record(record)
    engine = value["engine_metadata"]
    lines = [
        f"# {value['report_title']}", "",
        f"**Matter:** {value['matter_name']}",
        f"**Jurisdiction:** {value['jurisdiction']}",
        f"**Analysis provider:** {engine.get('provider')}",
        f"**Model:** {engine.get('model')}",
        f"**Engine source:** {engine.get('source')}",
        f"**Generated:** {value['generated_at_utc']}", "",
    ]
    if value.get("include_contents"):
        lines.extend(["## Contents", ""])
        lines.extend(f"{index}. {label}" for index, (_, label, _) in enumerate(_visible_sections(value), 1))
        lines.append("")
    for _, label, section_value in _visible_sections(value):
        section_value = _presentation_value(section_value, value.get("include_source_references", True))
        lines.extend([f"## {label}", ""])
        rows = list(_rows(section_value))
        if len(rows) == 1 and not rows[0][0]:
            lines.extend([rows[0][1], ""])
        else:
            for left, right in rows:
                lines.append(f"- **{left}:** {right}" if left else right)
            lines.append("")
    lines.extend([
        "## Important Professional Review Notice", "",
        "This document is AI-assisted lawyer reference material generated from supplied application state. "
        "It is not legal advice, a final legal opinion, or proof of compliance with a court or regulator's rules. "
        "Qualified counsel must verify the facts, original evidence, law, jurisdiction, deadlines, authorities, "
        "procedural requirements and final wording before reliance or external use.", "",
    ])
    return "\n".join(lines)


def build_professional_docx(record: Dict[str, Any]) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except Exception as exc:
        raise RuntimeError("Editable Word export requires python-docx. Install requirements_professional_reports.txt.") from exc

    value = normalise_professional_record(record)
    engine = value["engine_metadata"]
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    document.styles["Normal"].font.name = "Aptos"
    document.styles["Normal"].font.size = Pt(10.5)
    header = section.header.paragraphs[0]
    header.text = "AI Lawyer Opposition - Offline Professional Report"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = section.footer.paragraphs[0]
    footer.text = "AI-assisted lawyer reference material - professional review required"
    if value.get("include_page_numbers"):
        footer.add_run("  |  Page ")
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        footer._p.append(field)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(value["report_title"])
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(23, 32, 51)
    meta = document.add_table(rows=0, cols=2)
    meta.style = "Table Grid"
    for label, item in (
        ("Matter", value["matter_name"]), ("Jurisdiction", value["jurisdiction"]),
        ("Analysis provider", engine.get("provider")), ("Model", engine.get("model")),
        ("Engine source", engine.get("source")), ("Generated", value["generated_at_utc"]),
    ):
        cells = meta.add_row().cells
        cells[0].text = _text(label)
        cells[1].text = _text(item)
    visible = list(_visible_sections(value))
    if value.get("include_contents"):
        document.add_heading("Contents", level=1)
        for _, label, _ in visible:
            document.add_paragraph(label, style="List Number")
    for _, label, section_value in visible:
        section_value = _presentation_value(section_value, value.get("include_source_references", True))
        document.add_heading(label, level=1)
        rows = list(_rows(section_value))
        if len(rows) == 1 and not rows[0][0]:
            document.add_paragraph(rows[0][1])
            continue
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "No. / Field"
        table.rows[0].cells[1].text = "Details"
        for left, right in rows:
            cells = table.add_row().cells
            cells[0].text = left
            cells[1].text = right
    document.add_page_break()
    document.add_heading("Important Professional Review Notice", level=1)
    document.add_paragraph(build_professional_markdown(value).split("## Important Professional Review Notice", 1)[1].strip())
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def build_professional_pdf(record: Dict[str, Any]) -> bytes:
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Table, TableStyle
    except Exception as exc:
        raise RuntimeError("PDF export requires reportlab. Install requirements_professional_reports.txt.") from exc

    value = normalise_professional_record(record)
    engine = value["engine_metadata"]
    output = io.BytesIO()
    document = SimpleDocTemplate(output, pagesize=A4, rightMargin=18*mm, leftMargin=18*mm, topMargin=20*mm, bottomMargin=18*mm, title=value["report_title"])
    styles = getSampleStyleSheet()
    title = ParagraphStyle("OfflineTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=21, leading=26, alignment=TA_CENTER, textColor=colors.HexColor("#172033"))
    heading = ParagraphStyle("OfflineHeading", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=14, leading=18, textColor=colors.HexColor("#3157D5"), spaceBefore=10, spaceAfter=6)
    body = ParagraphStyle("OfflineBody", parent=styles["BodyText"], fontSize=9.8, leading=14, textColor=colors.HexColor("#172033"))
    story: List[Any] = [
        Paragraph(html.escape(value["report_title"]), title),
        Paragraph(
            f"<b>Matter:</b> {html.escape(value['matter_name'])}<br/>"
            f"<b>Jurisdiction:</b> {html.escape(value['jurisdiction'])}<br/>"
            f"<b>Analysis provider:</b> {html.escape(_text(engine.get('provider')))}<br/>"
            f"<b>Model:</b> {html.escape(_text(engine.get('model')))}<br/>"
            f"<b>Engine source:</b> {html.escape(_text(engine.get('source')))}<br/>"
            f"<b>Generated:</b> {html.escape(value['generated_at_utc'])}", body,
        ),
    ]
    visible = list(_visible_sections(value))
    if value.get("include_contents"):
        story.append(Paragraph("Contents", heading))
        story.append(Paragraph("<br/>".join(f"{i}. {html.escape(label)}" for i, (_, label, _) in enumerate(visible, 1)), body))
    for _, label, section_value in visible:
        section_value = _presentation_value(section_value, value.get("include_source_references", True))
        story.append(Paragraph(html.escape(label), heading))
        rows = list(_rows(section_value))
        if len(rows) == 1 and not rows[0][0]:
            story.append(Paragraph(html.escape(rows[0][1]).replace("\n", "<br/>"), body))
            continue
        data = [[Paragraph("No. / Field", body), Paragraph("Details", body)]]
        data.extend([[Paragraph(html.escape(left), body), Paragraph(html.escape(right).replace("\n", "<br/>"), body)] for left, right in rows])
        table = Table(data, colWidths=[30*mm, 138*mm], repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#F4F6FA")),
            ("GRID", (0,0), (-1,-1), 0.4, colors.HexColor("#DCE2EA")),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("LEFTPADDING", (0,0), (-1,-1), 5), ("RIGHTPADDING", (0,0), (-1,-1), 5),
            ("TOPPADDING", (0,0), (-1,-1), 4), ("BOTTOMPADDING", (0,0), (-1,-1), 4),
        ]))
        story.append(table)
    story.extend([
        PageBreak(), Paragraph("Important Professional Review Notice", heading),
        Paragraph(build_professional_markdown(value).split("## Important Professional Review Notice", 1)[1].strip(), body),
    ])

    def page(canvas, doc):
        canvas.saveState()
        width, _ = A4
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#687386"))
        canvas.drawString(18*mm, 10*mm, "AI-assisted lawyer reference material - professional review required")
        if value.get("include_page_numbers"):
            canvas.drawRightString(width-18*mm, 10*mm, f"Page {doc.page}")
        canvas.restoreState()

    document.build(story, onFirstPage=page, onLaterPages=page)
    return output.getvalue()


def record_as_json(record: Dict[str, Any]) -> str:
    return json.dumps(normalise_professional_record(record), ensure_ascii=False, indent=2)
